

from docx import Document
from docx.shared import Inches


def main():
    document = Document()   # 建立文档
    document.add_heading('Document Title',0)   # 设置title
    p = document.add_paragraph('A plain paragraph having some ')  # 添加段落
    p.add_run('bold').bold = True     # 设置字体加粗
    p.add_run(' and some ')  # 段落添加文字
    p.add_run('italic.').italic = True  # 设置字体斜体
    p.add_run('shadow').shadow = True

    document.add_heading('Heading, level 1', level=1)
    document.add_paragraph('Intense quote', style='Intense Quote')

    document.add_paragraph(
        'first item in unordered list', style='List Bullet'
    )
    document.add_paragraph(
        'first item in ordered list', style='List Number'
    )

    document.add_picture('img/erweima.jpg', width=Inches(1.25))  # 添加图片

    records = (
        (3, '101', 'Spam'),
        (7, '422', 'Eggs'),
        (4, '631', 'Spam, spam, eggs, and spam')
    )

    table = document.add_table(rows=1, cols=3)   # 表格填写
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Qty'
    hdr_cells[1].text = 'Id'
    hdr_cells[2].text = 'Desc'
    for qty, id, desc in records:
        row_cells = table.add_row().cells
        row_cells[0].text = str(qty)
        row_cells[1].text = id
        row_cells[2].text = desc

    document.add_page_break()   # 新添页码

    document.save('demo.docx')






if __name__ == '__main__':
    filename = "爬虫.docx"
    main()
